"""Word document converter implementation."""

import os
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Union

import structlog
import docx

from morag_core.interfaces.converter import (
    ConversionOptions,
    ConversionError,
)
from morag_core.models.document import Document

from .base import DocumentConverter

logger = structlog.get_logger(__name__)


class WordConverter(DocumentConverter):
    """Word document converter implementation."""

    def __init__(self):
        """Initialize Word converter."""
        super().__init__()
        self.supported_formats = {"word", "docx", "doc"}

    async def _extract_text(self, file_path: Path, document: Document, options: ConversionOptions) -> Document:
        """Extract text from Word document.

        Args:
            file_path: Path to Word file
            document: Document to update
            options: Conversion options

        Returns:
            Updated document

        Raises:
            ConversionError: If text extraction fails
        """
        try:
            # Open Word document
            doc = docx.Document(file_path)
            
            # Extract document properties
            core_properties = doc.core_properties
            document.metadata.title = core_properties.title
            document.metadata.author = core_properties.author
            document.metadata.created_at = core_properties.created
            document.metadata.modified_at = core_properties.modified
            
            # Extract text from paragraphs
            full_text = ""
            current_section = ""
            
            # Process headings and paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # Check if paragraph is a heading
                if paragraph.style.name.startswith('Heading'):
                    # If we have content in the current section, add it as a chunk
                    if options.chunking_strategy == "section" and current_section and full_text:
                        document.add_chunk(
                            content=current_section,
                            section=text,  # Use heading as section name
                        )
                        current_section = ""
                    
                    # Start a new section with the heading
                    current_section = text + "\n\n"
                else:
                    # Add paragraph to current section
                    current_section += text + "\n\n"
                
                # Add to full text
                full_text += text + "\n\n"
            
            # Add the last section as a chunk if needed
            if options.chunking_strategy == "section" and current_section:
                document.add_chunk(
                    content=current_section,
                    section="Last Section",
                )
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        full_text += row_text + "\n"
            
            # Set raw text
            document.raw_text = full_text
            
            # Estimate word count
            document.metadata.word_count = len(full_text.split())
            
            return document
            
        except Exception as e:
            logger.error(
                "Word document text extraction failed",
                error=str(e),
                error_type=e.__class__.__name__,
                file_path=str(file_path),
            )
            raise ConversionError(f"Failed to extract text from Word document: {str(e)}")